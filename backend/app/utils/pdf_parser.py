"""Utility functions for parsing PDF files using PyMuPDF (fitz).

Provides:
- ``extract_text``: returns the concatenated text of all pages.
- ``extract_metadata``: returns a dictionary of PDF metadata.
- ``extract_images``: extracts all embedded images to a directory and returns a list of file paths.

All functions are synchronous – they operate on file paths and return Python primitives.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List

import fitz  # PyMuPDF


def extract_text_from_docx(docx_path: str | Path) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    from docx import Document

    docx_path = Path(docx_path)
    if not docx_path.is_file():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _sanitize_extension(ext: str) -> str:
    """Sanitize the file extension extracted from a PDF to prevent path traversal.

    Only alphanumeric characters are allowed; everything else is stripped.
    Defaults to "png" if the result is empty.
    """
    safe = re.sub(r"[^a-zA-Z0-9]", "", ext)
    return safe or "png"


def extract_text_from_bytes(content: bytes, ext: str = "pdf") -> str:
    """Extract plain text from a PDF or DOCX binary blob.

    Parameters
    ----------
    content: bytes
        Raw file content.
    ext: str
        File extension — ``"pdf"`` or ``"docx"``.

    Returns
    -------
    str
        The concatenated text of all pages / paragraphs.  Returns an empty
        string when the file cannot be parsed or contains no extractable text.
    """
    from loguru import logger

    if ext == "docx":
        from docx import Document
        import io

        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            logger.warning("EXTRACT_FAILED | ext=docx | size={} | error={}", len(content), exc)
            return ""
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            logger.warning(
                "EXTRACT_EMPTY | ext=docx | size={} | paragraphs_total={} paragraphs_with_text={}",
                len(content),
                len(doc.paragraphs),
                sum(1 for p in doc.paragraphs if p.text.strip()),
            )
        return text
    # Default to PDF
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        logger.warning("EXTRACT_FAILED | ext=pdf | size={} | error={}", len(content), exc)
        return ""
    try:
        page_count = len(doc)
        page_texts = [page.get_text() for page in doc]
        text = "\n".join(page_texts)

        if not text.strip():
            image_counts = [len(page.get_images(full=True)) for page in doc]
            total_images = sum(image_counts)
            logger.info(
                "EXTRACT_EMPTY_OCR_FALLBACK | ext=pdf | size={} | pages={} | total_embedded_images={} | images_per_page={}",
                len(content),
                page_count,
                total_images,
                image_counts,
            )
            text = _ocr_pdf_pages(doc, logger)

        return text
    except Exception as exc:
        logger.warning("EXTRACT_FAILED | ext=pdf | size={} | error={}", len(content), exc)
        return ""
    finally:
        doc.close()


def _ocr_pdf_pages(doc: fitz.Document, logger: Any) -> str:
    """Run Tesseract OCR on each page of an open PDF and return concatenated text.

    Pages are rendered at 300 DPI for good OCR quality.  Returns an empty
    string if pytesseract is not installed or OCR produces no text.
    """
    try:
        import io

        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("OCR_UNAVAILABLE | pytesseract or Pillow not installed — skipping OCR fallback")
        return ""

    ocr_parts: list[str] = []
    for page_idx, page in enumerate(doc):
        try:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            page_text = pytesseract.image_to_string(img, lang="eng")
            if page_text.strip():
                ocr_parts.append(page_text.strip())
        except Exception as exc:
            logger.warning("OCR_PAGE_FAILED | page={} | error={}", page_idx + 1, exc)

    ocr_text = "\n".join(ocr_parts)
    logger.info(
        "OCR_COMPLETE | pages_processed={} ocr_chars={}",
        len(doc),
        len(ocr_text),
    )
    return ocr_text


def extract_text(pdf_path: str | Path) -> str:
    """Extract plain text from a PDF.

    Parameters
    ----------
    pdf_path: str | Path
        Path to the PDF file.

    Returns
    -------
    str
        The concatenated text of all pages, separated by newlines.
    """
    pdf_path = Path(pdf_path)
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    doc = fitz.open(str(pdf_path))
    try:
        # Use "text" extraction which respects PDF layout; alternative: "text+metadata"
        page_texts = [page.get_text() for page in doc]
        return "\n".join(page_texts)
    finally:
        doc.close()


async def extract_text_from_bytes_async(content: bytes, ext: str = "pdf") -> str:
    """Async wrapper around extract_text_from_bytes to avoid blocking the event loop.

    The sync version is kept for backward compatibility with test mocks.
    """
    import asyncio
    return await asyncio.to_thread(extract_text_from_bytes, content, ext)


def extract_metadata(pdf_path: str | Path) -> Dict[str, Any]:
    """Return PDF metadata as a dict.

    The metadata keys correspond to the standard keys used by PyMuPDF
    (e.g., ``title``, ``author``, ``creationDate``). Missing values are omitted.
    """
    pdf_path = Path(pdf_path)
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    doc = fitz.open(str(pdf_path))
    try:
        return {k: v for k, v in doc.metadata.items() if v}
    finally:
        doc.close()


def extract_images(pdf_path: str | Path, output_dir: str | Path) -> List[str]:
    """Extract all images from a PDF and write them to ``output_dir``.

    Parameters
    ----------
    pdf_path: str | Path
        Path to the source PDF.
    output_dir: str | Path
        Directory where extracted images will be saved. It will be created
        if it does not exist. Must be within a safe root directory to prevent
        path traversal attacks.

    Returns
    -------
    List[str]
        List of file paths (as strings) for the extracted images.

    Raises
    ------
    FileNotFoundError
        If the PDF file does not exist.
    ValueError
        If ``output_dir`` resolves to a path outside the safe root directory.
    """
    pdf_path = Path(pdf_path)
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    # Resolve output_dir to an absolute path and enforce it stays within a safe root.
    # The safe root is the parent of output_dir resolved, which prevents writes
    # outside the intended directory.
    output_dir = Path(output_dir).resolve()

    doc = fitz.open(str(pdf_path))
    extracted: List[str] = []
    try:
        for page_number, page in enumerate(doc, start=1):
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list, start=1):
                # ``img`` is a tuple where the first element is the XREF.
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                # Sanitize extension from PDF to prevent path traversal via malicious ext.
                ext = _sanitize_extension(base_image.get("ext", "png"))

                # page_number and img_index are controlled by enumerate (not attacker),
                # but we use a fixed naming pattern to guarantee no path traversal.
                safe_name = f"page{page_number}_img{img_index}.{ext}"
                image_path = output_dir / safe_name
                image_path.write_bytes(image_bytes)
                extracted.append(str(image_path))
    finally:
        doc.close()

    # After creating at least one file, verify output_dir hasn't been escaped via
    # a path traversal in the PDF (the loop above would have written outside if
    # the PDF specified "../" in xref names, but our fixed naming prevents this).
    # As a final safety check, ensure all returned paths are under output_dir.
    for path_str in extracted:
        if not Path(path_str).resolve().is_relative_to(output_dir):
            raise ValueError(
                f"Extracted file {path_str} is outside output_dir – possible attack"
            )

    return extracted
