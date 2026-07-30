"""Programmatic Base-Nova DOCX template builder using python-docx.

Builds a two-column resume with a left sidebar (skills, languages,
certifications) and main content area.  All styling is done via
``python-docx`` primitives — no XML injection, no pdf2docx dependency.

Layout specs (A4 page):
- Page margins: 0.5 in top/bottom, 0.6 in left/right
- Sidebar column: 2.2 in, gutter 0.25 in, main fills remaining width
- Font: Karla throughout — 20 pt name, 11 pt headers, 10.5 pt roles,
  9.5 pt body at 1.15 line-spacing
- Section headings: 12 pt space-before, 6 pt space-after, 11 pt bold
  uppercase with 0.5 pt letter-spacing and a #00FFF0 underline
- Bullets: 0.15 in left-indent, 2 pt space-after
- Sidebar: 8 pt gap label→content, 14 pt between sub-sections
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Inches, RGBColor, Emu

# ── brand constants ──────────────────────────────────────────────────
BRAND_COLOR = RGBColor(0x00, 0xFF, 0xF0)  # #00FFF0
FONT_FAMILY = "Karla"

# ── page geometry (A4) ───────────────────────────────────────────────
MARGIN_TOP = Inches(0.5)
MARGIN_BOTTOM = Inches(0.5)
MARGIN_LEFT = Inches(0.6)
MARGIN_RIGHT = Inches(0.6)
PAGE_WIDTH_IN = 8.27
SIDEBAR_WIDTH_IN = 2.2
GUTTER_IN = 0.25
_MAIN_WIDTH_IN = PAGE_WIDTH_IN - 0.6 - 0.6 - SIDEBAR_WIDTH_IN - GUTTER_IN

# ── typography sizes ─────────────────────────────────────────────────
NAME_SIZE = Pt(20)
HEADER_SIZE = Pt(11)
ROLE_SIZE = Pt(10.5)
BODY_SIZE = Pt(9.5)

# ── spacing ──────────────────────────────────────────────────────────
SECTION_SPACE_BEFORE = Pt(12)
SECTION_SPACE_AFTER = Pt(6)
BULLET_INDENT = Inches(0.15)
BULLET_SPACE_AFTER = Pt(2)
SIDEBAR_LABEL_GAP = Pt(8)
SIDEBAR_SECTION_GAP = Pt(14)


# ── helpers ──────────────────────────────────────────────────────────

def _remove_table_borders(table: Any) -> None:
    """Strip every border from a table element."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    # remove existing borders element if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _set_cell_width(cell: Any, width_in: float) -> None:
    """Pin a table-cell width and prevent Word from auto-fitting."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # explicit width
    tcW = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(width_in * 1440)}" w:type="dxa"/>'
    )
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)
    # no-wrap
    noWrap = parse_xml(f'<w:noWrap {nsdecls("w")} w:val="true"/>')
    existing_nw = tcPr.find(qn("w:noWrap"))
    if existing_nw is not None:
        tcPr.remove(existing_nw)
    tcPr.append(noWrap)


def _set_no_autofit(table: Any) -> None:
    """Disable autofit on the table so column widths are honoured."""
    tblPr = table._tbl.tblPr
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    existing = tblPr.find(qn("w:tblLayout"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)


def _add_section_heading(container: Any, text: str) -> None:
    """Append a styled section heading paragraph.

    11 pt Karla Bold, uppercase, letter-spacing +0.5 pt,
    12 pt space-before / 6 pt space-after, #00FFF0 underline.
    """
    para = container.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = SECTION_SPACE_BEFORE
    pf.space_after = SECTION_SPACE_AFTER
    pf.line_spacing = 1.0

    run = para.add_run(text.upper())
    run.font.name = FONT_FAMILY
    run.font.size = HEADER_SIZE
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x17, 0x27)  # near-black

    # letter-spacing via XML (w:spacing in hundredths of a point → +0.5pt = 50)
    rPr = run._r.get_or_add_rPr()
    spacing_el = parse_xml(f'<w:spacing {nsdecls("w")} w:val="50"/>')
    rPr.append(spacing_el)

    # underline in brand colour
    _add_underline(para, BRAND_COLOR)


def _add_underline(para: Any, color: RGBColor) -> None:
    """Draw a bottom-border line spanning the paragraph width (brand color)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1"'
        f'            w:color="{str(color)}"/>'
        f'</w:pBdr>'
    )
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def _set_cell_vertical_align(cell: Any, align: str = "top") -> None:
    """Set vertical alignment on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(vAlign)


def _add_role_paragraph(container: Any, role: str, duration: str) -> None:
    """Add an inline role + duration paragraph.  10.5 pt SemiBold."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15

    if role:
        run_role = para.add_run(role)
        run_role.font.name = FONT_FAMILY
        run_role.font.size = ROLE_SIZE
        run_role.font.bold = True

    if duration:
        if role:
            sep = para.add_run("  |  ")
            sep.font.name = FONT_FAMILY
            sep.font.size = BODY_SIZE
            sep.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_dur = para.add_run(duration)
        run_dur.font.name = FONT_FAMILY
        run_dur.font.size = BODY_SIZE
        run_dur.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_body_text(container: Any, text: str) -> None:
    """Append a 9.5 pt body paragraph with 1.15 line-spacing."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = BODY_SIZE


def _add_bullet(container: Any, text: str) -> None:
    """Append a single bullet point.  0.15 in indent, 2 pt space-after."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.left_indent = BULLET_INDENT
    pf.first_line_indent = Inches(-0.1)
    pf.space_before = Pt(0)
    pf.space_after = BULLET_SPACE_AFTER
    pf.line_spacing = 1.15

    # bullet character
    bullet_run = para.add_run("•  ")
    bullet_run.font.name = FONT_FAMILY
    bullet_run.font.size = BODY_SIZE

    # text
    text_run = para.add_run(text)
    text_run.font.name = FONT_FAMILY
    text_run.font.size = BODY_SIZE


def _add_sidebar_label(container: Any, text: str) -> None:
    """Small uppercase label in the sidebar. 8 pt Karla Bold."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = SIDEBAR_SECTION_GAP
    pf.space_after = SIDEBAR_LABEL_GAP
    pf.line_spacing = 1.0

    run = para.add_run(text.upper())
    run.font.name = FONT_FAMILY
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = BRAND_COLOR


def _add_sidebar_text(container: Any, text: str, bold: bool = False) -> None:
    """Body text in the sidebar. 9.5 pt Karla."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15

    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = BODY_SIZE
    run.font.bold = bold


def _add_sidebar_bullet(container: Any, text: str) -> None:
    """A bullet item in the sidebar."""
    para = container.add_paragraph()
    pf = para.paragraph_format
    pf.left_indent = Inches(0.1)
    pf.first_line_indent = Inches(-0.08)
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15

    bullet_run = para.add_run("•  ")
    bullet_run.font.name = FONT_FAMILY
    bullet_run.font.size = Pt(8.5)

    text_run = para.add_run(text)
    text_run.font.name = FONT_FAMILY
    text_run.font.size = Pt(9)


def _format_duration(start: str, end: str, current: bool = False) -> str:
    """Build a human-readable duration string."""
    if current:
        end_str = "Present"
    else:
        end_str = end or "Present"
    if start:
        return f"{start} – {end_str}"
    return end_str


# ── public API ───────────────────────────────────────────────────────

def build_base_nova_template() -> Document:  # type: ignore[type-arg]
    """Return a blank ``Document`` with Base-Nova page geometry applied."""
    doc: Document = Document()  # type: ignore[call-arg]
    section = doc.sections[0]
    section.page_width = Emu(int(PAGE_WIDTH_IN * 914400))
    section.page_height = Emu(int(11.69 * 914400))
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # Remove default paragraph spacing so we control it explicitly
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY  # type: ignore[union-attr]
    style.font.size = BODY_SIZE  # type: ignore[union-attr]
    pf = style.paragraph_format  # type: ignore[union-attr]
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15

    return doc


def make_resume_docx(resume_data: dict[str, Any], output_path: str | Path) -> None:
    """Build a complete Base-Nova resume and write it to *output_path*.

    Parameters
    ----------
    resume_data:
        Dictionary mirroring the frontend ``ResumeData`` shape.  Expected
        top-level keys: ``content`` (with ``contact``, ``summary``,
        ``experience``, ``education``, ``skills``, ``languages``,
        ``certifications``, ``projects``, ``awards``, ``interests``,
        ``references``, ``custom``) and ``sections`` (list of
        ``{type, title, visible}``).
    output_path:
        Destination ``.docx`` file path.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_base_nova_template()
    content: dict[str, Any] = resume_data.get("content", {})
    contact: dict[str, str] = content.get("contact", {})
    sections: list[dict[str, Any]] = resume_data.get("sections", [])

    # Build a visibility / ordering map
    visible_types: dict[str, str] = {}
    ordered_types: list[str] = []
    for sec in sections:
        if sec.get("visible", True):
            t = sec.get("type", "")
            visible_types[t] = sec.get("title", t.title())
            ordered_types.append(t)

    # ── full-width header ────────────────────────────────────────────
    _build_header(doc, contact)

    # ── split sidebar vs main sections ───────────────────────────────
    sidebar_types = {"skills", "languages", "certifications", "awards",
                     "interests", "references"}
    main_order = [t for t in ordered_types if t not in sidebar_types]
    side_order = [t for t in ordered_types if t in sidebar_types]

    # ── two-column table ─────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _set_no_autofit(table)
    table.alignment = 0  # LEFT

    sidebar_cell = table.cell(0, 0)
    main_cell = table.cell(0, 1)

    _set_cell_width(sidebar_cell, SIDEBAR_WIDTH_IN)
    _set_cell_width(main_cell, _MAIN_WIDTH_IN)
    _set_cell_vertical_align(sidebar_cell, "top")
    _set_cell_vertical_align(main_cell, "top")

    # remove default empty paragraph in each cell
    for cell in (sidebar_cell, main_cell):
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

    # ── populate sidebar ─────────────────────────────────────────────
    for sec_type in side_order:
        title = visible_types.get(sec_type, sec_type.title())
        _render_sidebar_section(sidebar_cell, sec_type, title, content)

    # ── populate main ────────────────────────────────────────────────
    for sec_type in main_order:
        title = visible_types.get(sec_type, sec_type.title())
        _render_main_section(main_cell, sec_type, title, content)

    doc.save(str(output_path))


def inject_into_docx(resume_data: dict[str, Any], output_path: str | Path) -> None:
    """Thin wrapper — delegates to :func:`make_resume_docx`.

    Preserves the original interface so callers (routers / services)
    do not need to change.
    """
    make_resume_docx(resume_data, output_path)


# ── private renderers ────────────────────────────────────────────────

def _build_header(doc: Document, contact: dict[str, str]) -> None:  # type: ignore[type-arg]
    """Name, title, and contact line at the top of the page."""
    name = contact.get("fullName", "")
    title = contact.get("title", "")

    if name:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        run = para.add_run(name)
        run.font.name = FONT_FAMILY
        run.font.size = NAME_SIZE
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x11, 0x17, 0x27)

    if title:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.0
        run = para.add_run(title)
        run.font.name = FONT_FAMILY
        run.font.size = ROLE_SIZE
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # contact line
    parts = [contact.get("email", ""), contact.get("phone", ""),
             contact.get("location", "")]
    links = []
    if contact.get("linkedin"):
        links.append(contact["linkedin"])
    if contact.get("github"):
        links.append(contact["github"])
    if contact.get("website"):
        links.append(contact["website"])
    parts.extend(links)
    contact_text = "  •  ".join(p for p in parts if p)

    if contact_text:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(8)
        pf.line_spacing = 1.0
        run = para.add_run(contact_text)
        run.font.name = FONT_FAMILY
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _render_main_section(
    cell: Any, sec_type: str, title: str, content: dict[str, Any]
) -> None:
    """Render a section into the right (main) column cell."""
    if sec_type == "summary":
        _render_summary(cell, content.get("summary", ""))
    elif sec_type == "experience":
        _render_experience(cell, content.get("experience", []), title)
    elif sec_type == "education":
        _render_education(cell, content.get("education", []), title)
    elif sec_type == "projects":
        _render_projects(cell, content.get("projects", []), title)


def _render_summary(container: Any, text: str) -> None:
    if not text:
        return
    _add_section_heading(container, "Professional Summary")
    _add_body_text(container, text)


def _render_experience(container: Any, items: list[dict], title: str) -> None:
    if not items:
        return
    _add_section_heading(container, title)
    for idx, item in enumerate(items):
        if idx > 0:
            # thin separator between entries
            sep = container.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(2)

        duration = _format_duration(
            item.get("startDate", ""), item.get("endDate", ""),
            item.get("current", False),
        )
        company = item.get("company", "")
        role_line = f"{item.get('role', '')}" + (f" — {company}" if company else "")
        loc = item.get("location", "")
        if loc:
            duration = f"{duration}  •  {loc}" if duration else loc

        _add_role_paragraph(container, role_line, duration)

        for bullet_text in item.get("bullets", []):
            if bullet_text.strip():
                _add_bullet(container, bullet_text.strip())


def _render_education(container: Any, items: list[dict], title: str) -> None:
    if not items:
        return
    _add_section_heading(container, title)
    for idx, item in enumerate(items):
        if idx > 0:
            sep = container.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(2)

        degree = item.get("degree", "")
        field = item.get("field", "")
        degree_str = " — ".join(p for p in [degree, field] if p)

        institution = item.get("institution", "")
        duration = _format_duration(
            item.get("startDate", ""), item.get("endDate", ""),
            item.get("current", False),
        )

        _add_role_paragraph(container, degree_str, duration)

        if institution:
            _add_body_text(container, institution)

        gpa = item.get("gpa", "")
        if gpa:
            _add_body_text(container, f"GPA: {gpa}")


def _render_projects(container: Any, items: list[dict], title: str) -> None:
    if not items:
        return
    _add_section_heading(container, title)
    for idx, item in enumerate(items):
        if idx > 0:
            sep = container.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(2)

        name = item.get("name", "")
        role = item.get("role", "")
        duration = _format_duration(
            item.get("startDate", ""), item.get("endDate", ""),
            item.get("current", False),
        )
        role_line = f"{name}" + (f" — {role}" if role else "")
        _add_role_paragraph(container, role_line, duration)

        for bullet_text in item.get("bullets", []):
            if bullet_text.strip():
                _add_bullet(container, bullet_text.strip())


def _render_sidebar_section(
    cell: Any, sec_type: str, title: str, content: dict[str, Any]
) -> None:
    """Render a section into the left (sidebar) column cell."""
    if sec_type == "skills":
        _render_skills_sidebar(cell, content.get("skills", []), title)
    elif sec_type == "languages":
        _render_languages_sidebar(cell, content.get("languages", []), title)
    elif sec_type == "certifications":
        _render_certs_sidebar(cell, content.get("certifications", []), title)
    elif sec_type == "awards":
        _render_awards_sidebar(cell, content.get("awards", []), title)
    elif sec_type == "interests":
        _render_interests_sidebar(cell, content.get("interests", []), title)
    elif sec_type == "references":
        _render_references_sidebar(cell, content.get("references", []), title)


def _render_skills_sidebar(
    container: Any, groups: list[dict], title: str
) -> None:
    if not groups:
        return
    _add_sidebar_label(container, title)
    for group in groups:
        name = group.get("name", "")
        skills = group.get("skills", [])
        if name and skills:
            _add_sidebar_text(container, f"{name}: {', '.join(skills)}", bold=False)
        elif skills:
            _add_sidebar_text(container, ", ".join(skills))


def _render_languages_sidebar(
    container: Any, items: list[dict], title: str
) -> None:
    if not items:
        return
    _add_sidebar_label(container, title)
    for item in items:
        name = item.get("name", "")
        prof = item.get("proficiency", "")
        if name:
            display = f"{name}" + (f" — {prof}" if prof else "")
            _add_sidebar_text(container, display)


def _render_certs_sidebar(
    container: Any, items: list[dict], title: str
) -> None:
    if not items:
        return
    _add_sidebar_label(container, title)
    for item in items:
        name = item.get("name", "")
        issuer = item.get("issuer", "")
        date = item.get("date", "")
        parts = [name]
        if issuer:
            parts.append(issuer)
        if date:
            parts.append(date)
        _add_sidebar_text(container, " • ".join(p for p in parts if p))


def _render_awards_sidebar(
    container: Any, items: list[dict], title: str
) -> None:
    if not items:
        return
    _add_sidebar_label(container, title)
    for item in items:
        name = item.get("name", "")
        issuer = item.get("issuer", "")
        date = item.get("date", "")
        parts = [p for p in [name, issuer, date] if p]
        _add_sidebar_text(container, " • ".join(parts))


def _render_interests_sidebar(
    container: Any, items: list[dict], title: str
) -> None:
    if not items:
        return
    _add_sidebar_label(container, title)
    names = [item.get("name", "") for item in items if item.get("name")]
    if names:
        _add_sidebar_text(container, ", ".join(names))


def _render_references_sidebar(
    container: Any, items: list[dict], title: str
) -> None:
    if not items:
        return
    _add_sidebar_label(container, title)
    for item in items:
        name = item.get("name", "")
        role = item.get("role", "")
        company = item.get("company", "")
        email = item.get("email", "")
        line = " • ".join(p for p in [name, f"{role}, {company}" if role and company else role or company, email] if p)
        _add_sidebar_text(container, line)
